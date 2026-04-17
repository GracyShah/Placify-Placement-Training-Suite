import json
import math
import os
import re
from datetime import datetime
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document
from flask import current_app
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
from werkzeug.utils import secure_filename

from data.test_catalog import TOPIC_TEST_CATALOG
from nlp.resume_parser import ResumeParser


ALLOWED_RESUME_EXTENSIONS = {".pdf", ".docx"}
ALLOWED_JD_EXTENSIONS = {".pdf", ".docx", ".doc"}
DEFAULT_JOB_DESCRIPTION = (
    "We are hiring a software developer with python, java, javascript, sql, git, "
    "react, flask, docker, machine learning, and problem solving skills."
)
TOPIC_TEST_MAP = {
    test["test_id"]: test
    for topic in TOPIC_TEST_CATALOG
    for test in topic["tests"]
}
TOPIC_QUESTION_MAP = {
    question["question_key"]: {
        "topic_name": test["topic_name"],
        "question_text": question["question_text"],
        "difficulty": question.get("difficulty", "Medium"),
    }
    for test in TOPIC_TEST_MAP.values()
    for question in test["questions"]
}


def current_time():
    return datetime.now(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")


def calculate_resume_score(resume_data, jd_match_percentage=0):
    ats_score = 0
    if resume_data.get("full_name"):
        ats_score += 15
    if resume_data.get("email"):
        ats_score += 10
    if resume_data.get("phone"):
        ats_score += 10
    if resume_data.get("education"):
        ats_score += 20
    if resume_data.get("skills"):
        ats_score += 20
    if resume_data.get("experience") or resume_data.get("projects"):
        ats_score += 25

    keyword_score = 0
    skills_text = " ".join([
        resume_data.get("skills", ""),
        resume_data.get("experience", ""),
        resume_data.get("projects", ""),
    ]).lower()
    important_keywords = ["python", "java", "javascript", "react", "sql", "database", "api", "git", "team", "project", "leadership", "communication", "problem solving", "agile", "development"]
    for keyword in important_keywords:
        if keyword in skills_text:
            keyword_score += 100 / len(important_keywords)

    total_length = sum(len(str(resume_data.get(key, ""))) for key in ["education", "skills", "experience", "projects", "certifications"])
    format_score = 0
    if total_length > 100:
        format_score += 40
    if total_length > 300:
        format_score += 30
    if total_length > 500:
        format_score += 30

    overall_score = (ats_score + keyword_score + format_score) / 3
    overall_score = (overall_score * 0.6) + (jd_match_percentage * 0.4)

    feedback = []
    if ats_score < 70:
        feedback.append("Add more details to key sections like education and experience.")
    if keyword_score < 50:
        feedback.append("Include more relevant technical skills and keywords.")
    if format_score < 60:
        feedback.append("Expand your resume with more detailed descriptions.")
    if overall_score >= 80:
        feedback.append("Excellent resume! Well structured and comprehensive.")
    elif overall_score >= 60:
        feedback.append("Good resume, but there's room for improvement.")
    else:
        feedback.append("Your resume needs significant enhancement.")

    return {
        "ats_score": round(ats_score, 2),
        "keyword_score": round(keyword_score, 2),
        "format_score": round(format_score, 2),
        "overall_score": round(overall_score, 2),
        "feedback": " ".join(feedback),
        "jd_match_percentage": round(jd_match_percentage, 2),
    }


def build_ai_recommendation_payload(performance_rows):
    if not performance_rows:
        return {
            "weak_sections": [],
            "improvement_areas": [],
            "practice_focus": "Take a few tests to unlock personalized performance recommendations.",
            "readiness_score": 0,
            "topic_accuracy": [],
            "difficulty_analysis": {},
            "repeated_mistakes": [],
            "recommended_next_tests": [],
            "latest_recommendations": [],
        }

    attempt_scores = {}
    topic_stats = {}
    difficulty_stats = {}
    category_stats = {}
    incorrect_categories = {}

    for row in performance_rows:
        question_meta = TOPIC_QUESTION_MAP.get(row["category_name"], {})
        topic_name = row["section_name"] or row["topic_name"] or question_meta.get("topic_name") or "General"
        question_text = question_meta.get("question_text", "")
        difficulty = row["difficulty"] or question_meta.get("difficulty") or "Medium"
        attempt_key = (row["source_type"], row["topic_name"], row["test_name"], row["completed_at"])
        attempt_scores[attempt_key] = row["score"] or 0

        topic_bucket = topic_stats.setdefault(topic_name, {"correct": 0, "total": 0})
        topic_bucket["total"] += 1
        if row["is_correct"]:
            topic_bucket["correct"] += 1

        difficulty_bucket = difficulty_stats.setdefault(difficulty, {"correct": 0, "total": 0})
        difficulty_bucket["total"] += 1
        if row["is_correct"]:
            difficulty_bucket["correct"] += 1

        raw_category = row["category_name"] or ""
        if row["source_type"] == "company" and raw_category in TOPIC_CATEGORY_PATTERNS:
            category_name = raw_category
        elif row["source_type"] == "company" and raw_category in {"Logical Reasoning", "Verbal Ability"}:
            category_name = raw_category
        else:
            category_name = infer_topic_category(topic_name, question_text or raw_category)
        category_bucket = category_stats.setdefault(category_name, {"correct": 0, "total": 0})
        category_bucket["total"] += 1
        if row["is_correct"]:
            category_bucket["correct"] += 1
        else:
            incorrect_categories[category_name] = incorrect_categories.get(category_name, 0) + 1

    overall_avg = sum(attempt_scores.values()) / len(attempt_scores) if attempt_scores else 0
    topic_accuracy = [
        {"topic": topic, "accuracy": round((stats["correct"] / stats["total"]) * 100, 2), "questions_seen": stats["total"]}
        for topic, stats in sorted(topic_stats.items())
        if stats["total"]
    ]
    topic_accuracy.sort(key=lambda item: item["accuracy"])
    weak_sections = [item["topic"] for item in topic_accuracy if item["accuracy"] < 60][:4]

    difficulty_analysis = {
        level: round((stats["correct"] / stats["total"]) * 100, 2)
        for level, stats in difficulty_stats.items()
        if stats["total"]
    }
    repeated_mistakes = [name for name, count in sorted(incorrect_categories.items(), key=lambda item: item[1], reverse=True) if count >= 2][:5]

    recommendations = []
    for item in topic_accuracy[:3]:
        if item["accuracy"] < 60:
            recommendations.append(f"Your overall accuracy in {item['topic']} is only {item['accuracy']}%.")

    easy_accuracy = difficulty_analysis.get("Easy")
    medium_accuracy = difficulty_analysis.get("Medium")
    hard_accuracy = difficulty_analysis.get("Hard")
    if easy_accuracy is not None and medium_accuracy is not None and medium_accuracy + 12 < easy_accuracy:
        recommendations.append("Performance drops significantly in Medium-difficulty questions.")
    if medium_accuracy is not None and hard_accuracy is not None and hard_accuracy + 12 < medium_accuracy:
        recommendations.append("Hard-difficulty accuracy is trailing behind your medium sets.")
    if repeated_mistakes:
        recommendations.append(f"Repeated mistakes are appearing in {', '.join(repeated_mistakes[:3])}.")

    improvement_areas = []
    for section_name in weak_sections:
        if section_name == "Aptitude":
            improvement_areas.append("Practice more quantitative problems, especially percentages, averages, and time-based arithmetic.")
        elif section_name in {"Logical", "Logical Reasoning"}:
            improvement_areas.append("Work on pattern recognition, coding-decoding, and arrangement puzzles.")
        elif section_name in {"Verbal", "Verbal Ability"}:
            improvement_areas.append("Revise grammar rules, vocabulary usage, and sentence correction patterns.")
        elif section_name == "DSA Basics":
            improvement_areas.append("Strengthen core data structures, searching, sorting, and complexity analysis.")
        elif section_name == "SQL":
            improvement_areas.append("Revise joins, aggregation, filtering clauses, and key constraints.")
        else:
            improvement_areas.append(f"Spend more practice time improving {section_name}.")

    recommended_next_tests = [f"{section} next practice set" for section in weak_sections[:3]]
    if overall_avg < 50:
        practice_focus = "Focus on fundamentals first, then move back to timed mixed tests."
    elif overall_avg < 70:
        practice_focus = "You are progressing well. Prioritize weak topics and medium-difficulty consistency."
    else:
        practice_focus = "You are close to demo-ready placement performance. Maintain consistency and sharpen weak patterns."

    return {
        "weak_sections": weak_sections,
        "improvement_areas": improvement_areas[:5],
        "practice_focus": practice_focus,
        "readiness_score": round(min(overall_avg + 10, 100), 2),
        "topic_accuracy": topic_accuracy,
        "difficulty_analysis": difficulty_analysis,
        "repeated_mistakes": repeated_mistakes,
        "recommended_next_tests": recommended_next_tests,
        "latest_recommendations": recommendations[:5],
    }


def count_projects(text):
    return len(re.findall(r"\bproject[s]?\b", (text or "").lower()))


def detect_resume_type(resume_text):
    text = (resume_text or "").lower()
    project_count = count_projects(text)
    experience_entries = len(re.findall(r"\bexperience\b|\bworked\b|\bemployment\b|\bdeveloper\b", text))
    internship_hits = len(re.findall(r"\binternship\b|\bintern\b|\btrainee\b", text))

    years = []
    years.extend(float(match) for match in re.findall(r"(\d+(?:\.\d+)?)\+?\s+years?", text))
    years.extend(float(match) for match in re.findall(r"experience\s+of\s+(\d+(?:\.\d+)?)", text))
    max_years = max(years) if years else 0.0

    fresher_score = 0.0
    experienced_score = 0.0
    if max_years >= 2:
        experienced_score += 0.55
    elif max_years >= 1:
        experienced_score += 0.35
    else:
        fresher_score += 0.25
    if internship_hits:
        fresher_score += min(0.25, internship_hits * 0.08)
    if project_count >= max(experience_entries, 1):
        fresher_score += 0.25
    else:
        experienced_score += 0.2
    if experience_entries >= 2:
        experienced_score += 0.25
    elif experience_entries == 1:
        experienced_score += 0.1
    else:
        fresher_score += 0.15

    total_score = fresher_score + experienced_score
    if total_score == 0:
        return {"resume_type": "Fresher", "confidence_score": 0.5}

    fresher_confidence = fresher_score / total_score
    experienced_confidence = experienced_score / total_score
    if experienced_confidence > fresher_confidence:
        return {"resume_type": "Experienced professional", "confidence_score": round(experienced_confidence, 2)}
    return {"resume_type": "Fresher", "confidence_score": round(fresher_confidence, 2)}


def split_items(text):
    return [item.strip() for item in re.split(r"[\n,]+", text or "") if item.strip()]


def extract_capabilities(text, skills=None):
    normalized_text = (text or "").lower()
    extracted_skills = list(skills or [])

    technical_keywords = {"api", "backend", "frontend", "database", "algorithms", "data structures", "machine learning", "deep learning", "testing", "automation", "cloud"}
    domain_keywords = {"finance", "banking", "healthcare", "ecommerce", "education", "analytics", "nlp", "computer vision", "cybersecurity", "devops", "web development"}
    tool_keywords = {"git", "github", "docker", "kubernetes", "aws", "azure", "google cloud", "flask", "django", "fastapi", "react", "node.js", "express.js", "tensorflow", "pytorch", "pandas", "numpy", "scikit-learn"}

    technical_skills = sorted({skill for skill in extracted_skills if skill in technical_keywords or skill in tool_keywords} | {keyword for keyword in technical_keywords if keyword in normalized_text})
    domain_knowledge = sorted({skill for skill in extracted_skills if skill in domain_keywords} | {keyword for keyword in domain_keywords if keyword in normalized_text})
    tools_and_technologies = sorted({skill for skill in extracted_skills if skill in tool_keywords} | {keyword for keyword in tool_keywords if keyword in normalized_text})

    applied_work = []
    for label, patterns in {
        "projects": [r"\bproject[s]?\b", r"\bbuilt\b", r"\bdeveloped\b", r"\bcreated\b"],
        "experience": [r"\bexperience\b", r"\bworked\b", r"\bemployment\b", r"\bengineer\b"],
        "internships": [r"\binternship\b", r"\bintern\b", r"\btrainee\b"],
    }.items():
        if any(re.search(pattern, normalized_text) for pattern in patterns):
            applied_work.append(label)

    return {
        "technical_skills": technical_skills,
        "applied_work": applied_work,
        "domain_knowledge": domain_knowledge,
        "tools_and_technologies": tools_and_technologies,
    }


def compute_overlap_score(resume_values, jd_values):
    resume_set = {value.strip().lower() for value in resume_values if value and value.strip()}
    jd_set = {value.strip().lower() for value in jd_values if value and value.strip()}
    if not jd_set:
        return 100.0
    return round((len(resume_set & jd_set) / len(jd_set)) * 100, 2)


def get_capability_weights(resume_type):
    if resume_type == "Experienced professional":
        return {"technical_skills": 0.3, "applied_work": 0.35, "domain_knowledge": 0.2, "tools_and_technologies": 0.15}
    return {"technical_skills": 0.35, "applied_work": 0.2, "domain_knowledge": 0.15, "tools_and_technologies": 0.3}


def compute_capability_match_scores(resume_text, job_description_text, resume_skills, job_skills, resume_type):
    resume_capabilities = extract_capabilities(resume_text, resume_skills)
    jd_capabilities = extract_capabilities(job_description_text, job_skills)
    weights = get_capability_weights(resume_type)
    category_scores = {key: compute_overlap_score(resume_capabilities[key], jd_capabilities[key]) for key in weights}
    role_fit_score = round(sum(category_scores[key] * weights[key] for key in weights), 2)
    skill_readiness_score = round((category_scores["technical_skills"] * 0.6) + (category_scores["tools_and_technologies"] * 0.25) + (category_scores["domain_knowledge"] * 0.15), 2)
    return {
        "resume_capabilities": resume_capabilities,
        "jd_capabilities": jd_capabilities,
        "category_scores": category_scores,
        "role_fit_score": role_fit_score,
        "skill_readiness_score": skill_readiness_score,
        "learning_gap_score": round(max(0.0, 100.0 - role_fit_score), 2),
    }


def infer_skill_evidence(skill, resume_text):
    normalized_text = (resume_text or "").lower()
    normalized_skill = skill.strip().lower()
    evidence = {"projects": [], "experience": [], "tools_usage": []}
    project_patterns = [rf"project[s]?[^\n.]*{re.escape(normalized_skill)}", rf"{re.escape(normalized_skill)}[^\n.]*project[s]?"]
    experience_patterns = [rf"experience[^\n.]*{re.escape(normalized_skill)}", rf"worked[^\n.]*{re.escape(normalized_skill)}", rf"{re.escape(normalized_skill)}[^\n.]*experience"]
    tool_patterns = [rf"using[^\n.]*{re.escape(normalized_skill)}", rf"with[^\n.]*{re.escape(normalized_skill)}", rf"tool[s]?[^\n.]*{re.escape(normalized_skill)}"]
    for pattern in project_patterns:
        evidence["projects"].extend(re.findall(pattern, normalized_text))
    for pattern in experience_patterns:
        evidence["experience"].extend(re.findall(pattern, normalized_text))
    for pattern in tool_patterns:
        evidence["tools_usage"].extend(re.findall(pattern, normalized_text))
    return {key: list(dict.fromkeys(value)) for key, value in evidence.items()}


def get_skill_confidence_label(score):
    if score >= 75:
        return "High"
    if score >= 45:
        return "Medium"
    return "Low"


def score_skills_with_evidence(skills, resume_text):
    confidence = []
    for skill in skills:
        evidence = infer_skill_evidence(skill, resume_text)
        score = 20.0
        if evidence["projects"]:
            score += 35.0
        if evidence["experience"]:
            score += 35.0
        if evidence["tools_usage"]:
            score += 20.0
        if not any(evidence.values()):
            score -= 15.0
        score = max(5.0, min(score, 100.0))
        confidence.append({"skill": skill, "confidence_score": round(score, 2), "confidence_level": get_skill_confidence_label(score), "evidence": evidence})
    return confidence


def build_tailored_resume(data, job_description_text):
    from nlp.skill_extractor import SkillExtractor
    from nlp.text_preprocessing import TextPreprocessor

    preprocessor = TextPreprocessor()
    extractor = SkillExtractor()
    skills_input = data.get("skills", "")
    experience = data.get("experience", "")
    projects = data.get("projects", "")
    certifications = data.get("certifications", "")
    education = data.get("education", "")
    target_company = (data.get("target_company") or "the target company").strip()
    target_role = (data.get("target_role") or "the target role").strip()
    job_description_text = (job_description_text or DEFAULT_JOB_DESCRIPTION).strip()

    combined_text = " ".join([skills_input, experience, projects, certifications, education])
    user_skills = extractor.extract_skills(preprocessor.preprocess(combined_text))
    job_skills = extractor.extract_skills(preprocessor.preprocess(job_description_text))
    merged_skills = list(dict.fromkeys(user_skills + job_skills))
    ats_keywords = ", ".join(job_skills) if job_skills else "python, sql, communication"
    summary = (
        f"Targeting the {target_role} opportunity at {target_company}. "
        f"Candidate brings experience across {', '.join(merged_skills[:8]) or 'software development'} "
        f"with a resume aligned to the job description requirements."
    )
    sections = [
        ("Full Name", data.get("full_name", "")),
        ("Contact", f"Email: {data.get('email', '')}\nPhone: {data.get('phone', '')}"),
        ("Professional Summary", summary),
        ("Core Skills", ", ".join(merged_skills)),
        ("Professional Experience", experience or "Add relevant internships, work experience, or leadership roles."),
        ("Projects", projects or "Add project details with measurable outcomes and tools used."),
        ("Education", education),
        ("Certifications", certifications or "Add certifications relevant to the target role."),
        ("Target Company", target_company),
        ("Target Role", target_role),
        ("Job Description Alignment", job_description_text),
        ("ATS Keywords", ats_keywords),
    ]
    resume_text = "\n\n".join(f"{title}\n{content}".strip() for title, content in sections if content)
    return {"sections": sections, "resume_text": resume_text, "match_percentage": 100.0 if job_skills else 90.0}


def generate_resume_files(user_id, sections):
    output_dir = Path(current_app.config["DOWNLOAD_FOLDER"])
    output_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = output_dir / f"resume_{user_id}.pdf"
    docx_path = output_dir / f"resume_{user_id}.docx"
    generate_resume_pdf(pdf_path, sections)
    generate_resume_docx(docx_path, sections)
    return {"pdf": str(pdf_path), "doc": str(docx_path)}


def generate_resume_pdf(file_path, sections):
    pdf = canvas.Canvas(str(file_path), pagesize=LETTER)
    width, height = LETTER
    x_margin = 50
    y = height - 50
    for title, content in sections:
        if not content:
            continue
        if y < 100:
            pdf.showPage()
            y = height - 50
        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawString(x_margin, y, title)
        y -= 18
        pdf.setFont("Helvetica", 10)
        for line in str(content).splitlines():
            for wrapped_line in wrap_text(line, 90):
                if y < 60:
                    pdf.showPage()
                    y = height - 50
                    pdf.setFont("Helvetica", 10)
                pdf.drawString(x_margin, y, wrapped_line)
                y -= 14
        y -= 10
    pdf.save()


def generate_resume_docx(file_path, sections):
    document = Document()
    for index, (title, content) in enumerate(sections):
        if not content:
            continue
        if index == 0:
            document.add_heading(str(content), level=0)
            continue
        document.add_heading(title, level=1)
        for line in str(content).splitlines():
            document.add_paragraph(line)
    document.save(str(file_path))


def wrap_text(text, max_length):
    words = text.split()
    lines = []
    current_line = []
    for word in words:
        trial_line = " ".join(current_line + [word])
        if len(trial_line) <= max_length:
            current_line.append(word)
        else:
            lines.append(" ".join(current_line))
            current_line = [word]
    if current_line:
        lines.append(" ".join(current_line))
    return lines or [""]


def save_uploaded_file(uploaded_file):
    upload_dir = Path(current_app.config["UPLOAD_FOLDER"])
    upload_dir.mkdir(parents=True, exist_ok=True)
    filename = secure_filename(uploaded_file.filename)
    file_path = upload_dir / filename
    uploaded_file.save(file_path)
    return file_path


def extract_uploaded_document_text(uploaded_file, allowed_extensions):
    file_extension = os.path.splitext(uploaded_file.filename)[1].lower()
    if file_extension not in allowed_extensions:
        return None, file_extension
    file_path = save_uploaded_file(uploaded_file)
    if file_extension == ".doc":
        with open(file_path, "rb") as file_handle:
            return file_handle.read().decode("latin-1", errors="ignore"), file_extension
    return ResumeParser.extract_text(file_path), file_extension


def should_use_openai_resume_ai():
    return os.environ.get("ENABLE_OPENAI_RESUME_AI", "false").strip().lower() in {"1", "true", "yes", "on"}


def get_openai_client():
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("OPENAI_API_KEY is not set")
    try:
        from openai import OpenAI
    except ImportError as error:
        raise ValueError("OpenAI SDK is not installed") from error
    return OpenAI(api_key=api_key)


def extract_json_object(text):
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        start = text.find("{")
        end = text.rfind("}")
        if start == -1 or end == -1 or end <= start:
            raise ValueError("AI response was not valid JSON")
        return json.loads(text[start : end + 1])


def normalize_suggestions_payload(payload):
    return {
        "top_strengths": split_items(payload.get("top_strengths", ""))[:3],
        "priority_improvements": split_items(payload.get("priority_improvements", ""))[:5],
        "missing_keywords": split_items(payload.get("missing_keywords", ""))[:8],
        "rewritten_summary": (payload.get("rewritten_summary") or "").strip(),
        "bullet_improvements": split_items(payload.get("bullet_improvements", ""))[:5],
    }


def tokenize_for_tfidf(text):
    tokens = re.findall(r"[a-z0-9+#.]{2,}", (text or "").lower())
    stop_words = {
        "about", "after", "again", "along", "also", "an", "and", "are", "as", "at", "be", "because", "been",
        "before", "being", "between", "both", "build", "built", "can", "candidate", "consider", "currently", "data",
        "demo", "description", "developer", "developers", "do", "for", "from", "have", "help", "high", "if", "in",
        "into", "is", "it", "job", "knowledge", "low", "match", "more", "of", "on", "or", "our", "role", "should",
        "skills", "software", "strong", "team", "than", "that", "the", "their", "them", "then", "these", "this",
        "through", "to", "using", "very", "we", "well", "with", "work", "worked", "your",
    }
    return [token for token in tokens if token not in stop_words and len(token) > 2]


def build_tfidf_vector(tokens, document_sets):
    vector = {}
    total_terms = len(tokens) or 1
    vocabulary = set(tokens)
    for term in vocabulary:
        tf = tokens.count(term) / total_terms
        doc_frequency = sum(1 for doc in document_sets if term in doc)
        idf = math.log((1 + len(document_sets)) / (1 + doc_frequency)) + 1
        vector[term] = tf * idf
    return vector


def cosine_similarity_score(vector_a, vector_b):
    if not vector_a or not vector_b:
        return 0.0
    shared_terms = set(vector_a) | set(vector_b)
    dot_product = sum(vector_a.get(term, 0.0) * vector_b.get(term, 0.0) for term in shared_terms)
    magnitude_a = math.sqrt(sum(value * value for value in vector_a.values()))
    magnitude_b = math.sqrt(sum(value * value for value in vector_b.values()))
    if magnitude_a == 0 or magnitude_b == 0:
        return 0.0
    return round((dot_product / (magnitude_a * magnitude_b)) * 100, 2)


def extract_ranked_keywords(cleaned_text, max_keywords=12):
    tokens = tokenize_for_tfidf(cleaned_text)
    ranked = []
    seen = set()
    for token in tokens:
        if token not in seen:
            ranked.append(token)
            seen.add(token)
        if len(ranked) >= max_keywords:
            break
    return ranked


def estimate_shortlisting_probability(match_percentage):
    if match_percentage >= 80:
        return "High shortlisting probability"
    if match_percentage >= 60:
        return "Moderate shortlisting probability"
    if match_percentage >= 40:
        return "Needs improvement for shortlisting"
    return "Low shortlisting probability until resume alignment improves"


def build_section_improvements(missing_keywords, resume_text, match_percentage):
    lowered_resume = (resume_text or "").lower()
    improvements = []
    if missing_keywords:
        top_missing = ", ".join(missing_keywords[:3])
        improvements.append({
            "section": "Skills",
            "improvement": f"Add or strengthen evidence for {top_missing} so recruiters can quickly confirm role alignment.",
        })
    if "project" not in lowered_resume and missing_keywords:
        improvements.append({
            "section": "Projects",
            "improvement": f"Consider adding projects or experience related to {missing_keywords[0]} with outcomes, tools, and measurable impact.",
        })
    elif count_projects(resume_text) <= 1:
        improvements.append({
            "section": "Projects",
            "improvement": "Resume length appears short on project evidence; expand project descriptions with measurable impact, ownership, and tech stack details.",
        })
    if match_percentage < 55:
        improvements.append({
            "section": "Experience",
            "improvement": "Rephrase resume bullets using job description keywords for better ATS matching and clearer relevance to the target role.",
        })
    if "summary" not in lowered_resume and "objective" not in lowered_resume:
        improvements.append({
            "section": "Summary",
            "improvement": "Add a short professional summary that mirrors the role focus, core stack, and strongest achievements from your resume.",
        })
    return improvements[:4]


def build_ats_tips(missing_keywords, match_percentage, resume_text):
    tips = []
    if missing_keywords:
        tips.append(f"Use exact JD keywords such as {', '.join(missing_keywords[:3])} in relevant skills, project, or experience bullets.")
    if match_percentage < 60:
        tips.append("Your resume match score is low; focus on aligning technical skills, frameworks, and delivery outcomes with the role requirements.")
    if len((resume_text or '').split()) < 220:
        tips.append("Resume length appears short; expand project descriptions with measurable impact, scale, and ownership.")
    tips.append("Keep role-specific keywords close to achievements so ATS systems and recruiters see both skill and evidence together.")
    return tips[:4]


def generate_local_resume_ai_suggestions(resume_text, job_description_text, fallback_reason=None):
    from nlp.skill_extractor import SkillExtractor
    from nlp.text_preprocessing import TextPreprocessor

    preprocessor = TextPreprocessor()
    extractor = SkillExtractor()
    cleaned_resume = preprocessor.preprocess(resume_text)
    cleaned_job = preprocessor.preprocess(job_description_text)

    resume_tokens = tokenize_for_tfidf(cleaned_resume)
    job_tokens = tokenize_for_tfidf(cleaned_job)
    document_sets = [set(resume_tokens), set(job_tokens)]
    resume_vector = build_tfidf_vector(resume_tokens, document_sets)
    job_vector = build_tfidf_vector(job_tokens, document_sets)
    similarity_score = cosine_similarity_score(resume_vector, job_vector)

    resume_skills = [skill.lower() for skill in extractor.extract_skills(cleaned_resume)]
    job_skills = [skill.lower() for skill in extractor.extract_skills(cleaned_job)]
    jd_keywords = list(dict.fromkeys(job_skills + extract_ranked_keywords(cleaned_job)))
    resume_keywords = set(resume_skills + extract_ranked_keywords(cleaned_resume, max_keywords=18))
    missing_keywords = [keyword for keyword in jd_keywords if keyword not in resume_keywords][:8]
    weak_keywords = [keyword for keyword in jd_keywords if cleaned_resume.count(keyword) == 1][:5]

    strengths = []
    overlap_keywords = [keyword for keyword in jd_keywords if keyword in resume_keywords][:5]
    if overlap_keywords:
        strengths.append(f"Strong overlap detected for {', '.join(overlap_keywords[:3])}.")
    if similarity_score >= 65:
        strengths.append("Resume language is already reasonably aligned with the job description.")
    if count_projects(resume_text) >= 2:
        strengths.append("Project coverage gives you useful proof points for recruiter screening.")
    if not strengths:
        strengths.append("The resume has a baseline foundation, but stronger role alignment is needed for better screening outcomes.")

    section_improvements = build_section_improvements(missing_keywords or weak_keywords, resume_text, similarity_score)
    ats_tips = build_ats_tips(missing_keywords or weak_keywords, similarity_score, resume_text)
    priority_improvements = [item["improvement"] for item in section_improvements]
    if weak_keywords:
        priority_improvements.append(f"Some important JD terms are weakly represented: {', '.join(weak_keywords[:3])}. Add stronger evidence around them.")

    rewritten_summary = (
        f"Role-aligned candidate with experience in {', '.join((overlap_keywords or jd_keywords)[:4]) or 'software development'}, "
        f"focused on delivering measurable outcomes and stronger alignment to the target job requirements."
    )

    notice = None
    if fallback_reason or not should_use_openai_resume_ai():
        notice = "Advanced AI suggestions currently unavailable. Showing intelligent NLP-based recommendations."

    return {
        "source": "nlp",
        "notice": notice,
        "match_percentage": similarity_score,
        "skill_gaps": missing_keywords[:6],
        "missing_keywords": missing_keywords[:8],
        "section_wise_improvements": section_improvements,
        "ats_optimization_tips": ats_tips,
        "shortlisting_probability_estimate": estimate_shortlisting_probability(similarity_score),
        "top_strengths": strengths[:3],
        "priority_improvements": priority_improvements[:5],
        "rewritten_summary": rewritten_summary,
        "bullet_improvements": ats_tips[:3],
    }


def generate_resume_ai_suggestions(resume_text, job_description_text):
    if not should_use_openai_resume_ai():
        return generate_local_resume_ai_suggestions(resume_text, job_description_text)

    try:
        client = get_openai_client()
        response = client.responses.create(
            model=current_app.config["OPENAI_MODEL"],
            instructions="You are an expert resume coach. Return valid JSON with the keys top_strengths, priority_improvements, missing_keywords, rewritten_summary, and bullet_improvements.",
            input=[{"role": "user", "content": [{"type": "input_text", "text": f"Analyze this resume against the job description and suggest improvements.\n\nResume Text:\n{resume_text}\n\nJob Description Text:\n{job_description_text}"}]}],
        )
        normalized = normalize_suggestions_payload(extract_json_object(response.output_text))
        return {
            "source": "openai",
            "notice": None,
            "match_percentage": 0,
            "skill_gaps": normalized["missing_keywords"],
            "missing_keywords": normalized["missing_keywords"],
            "section_wise_improvements": [{"section": "Resume", "improvement": item} for item in normalized["priority_improvements"]],
            "ats_optimization_tips": normalized["bullet_improvements"],
            "shortlisting_probability_estimate": "AI-generated estimate available in narrative guidance",
            **normalized,
        }
    except Exception:
        return generate_local_resume_ai_suggestions(resume_text, job_description_text, fallback_reason="openai_unavailable")


def generate_local_resume_ai_chat_reply(resume_text, job_description_text, suggestions, question, chat_history):
    question_lower = (question or "").lower()
    missing_keywords = suggestions.get("missing_keywords") or []
    section_improvements = suggestions.get("section_wise_improvements") or []
    ats_tips = suggestions.get("ats_optimization_tips") or []
    match_percentage = suggestions.get("match_percentage") or 0

    if "skill" in question_lower or "missing" in question_lower:
        if missing_keywords:
            return f"The clearest skill gaps are {', '.join(missing_keywords[:4])}. Add them only where you have real evidence, ideally inside project or experience bullets tied to outcomes."
        return "Your resume already covers most of the visible JD keywords, so the next improvement is making those skills easier to spot through stronger bullet wording and measurable outcomes."
    if "experience" in question_lower or "project" in question_lower:
        project_guidance = next((item["improvement"] for item in section_improvements if item.get("section") in {"Experience", "Projects"}), None)
        return project_guidance or "Strengthen experience bullets by naming the stack, the problem solved, your ownership, and a measurable result such as speed, scale, or accuracy improvement."
    if "ats" in question_lower or "keyword" in question_lower:
        return " ".join(ats_tips[:2]) if ats_tips else "Mirror the job description language in your headings and bullets, especially around tools, frameworks, and delivery outcomes."
    if "score" in question_lower or "match" in question_lower or "chance" in question_lower:
        return f"Your current local match score is {round(match_percentage, 2)}%. The fastest way to improve it is to close the keyword gaps, strengthen project evidence, and rephrase bullets so they reflect the target role vocabulary."
    if "summary" in question_lower:
        return suggestions.get("rewritten_summary") or "Write a short summary that names your target role, strongest stack, and one or two outcome-focused strengths relevant to the JD."
    if missing_keywords:
        return f"A strong next step is to improve evidence around {missing_keywords[0]} and related JD language. Update the relevant skills, project, and experience bullets so the match is easier for ATS and recruiters to detect."
    return "Your resume is on the right track. Focus on clearer role alignment, measurable impact in bullets, and stronger keyword placement in the sections most relevant to the job description."


def generate_resume_ai_chat_reply(resume_text, job_description_text, suggestions, question, chat_history):
    if suggestions.get("source") == "nlp" or not should_use_openai_resume_ai():
        return generate_local_resume_ai_chat_reply(resume_text, job_description_text, suggestions, question, chat_history)

    try:
        client = get_openai_client()
        messages = [{"role": "user", "content": [{"type": "input_text", "text": f"Use the following context to answer follow-up resume improvement questions.\n\nResume Text:\n{resume_text}\n\nJob Description Text:\n{job_description_text}\n\nCurrent Suggestions:\n{json.dumps(suggestions)}"}]}]
        for message in chat_history[-6:]:
            role = message.get("role")
            content = (message.get("content") or "").strip()
            if role in {"user", "assistant"} and content:
                messages.append({"role": role, "content": [{"type": "input_text", "text": content}]})
        messages.append({"role": "user", "content": [{"type": "input_text", "text": question}]})
        response = client.responses.create(
            model=current_app.config["OPENAI_MODEL"],
            instructions="You are a resume coach. Answer follow-up questions using the supplied resume, job description, and prior suggestions. Keep responses clear and actionable.",
            input=messages,
        )
        return response.output_text.strip()
    except Exception:
        return generate_local_resume_ai_chat_reply(resume_text, job_description_text, suggestions, question, chat_history)


TOPIC_CATEGORY_PATTERNS = {
    "Aptitude": {
        "Time & Work": ["work", "pipe", "tank", "days", "finish"],
        "Percentages": ["percent", "%", "discount", "interest"],
        "Speed & Distance": ["train", "speed", "distance", "km", "hours"],
        "Average & Ratio": ["average", "ratio", "class", "boys", "girls"],
    },
    "Logical": {
        "Series & Patterns": ["series", "next", "missing term"],
        "Coding-Decoding": ["code", "coded", "written as"],
        "Directions & Arrangements": ["east", "west", "north", "south", "row"],
        "Syllogisms & Relations": ["all", "some", "related", "conclusion", "photograph"],
    },
    "Logical Reasoning": {
        "Series & Patterns": ["series", "next", "missing term"],
        "Coding-Decoding": ["code", "coded", "written as"],
        "Directions & Arrangements": ["east", "west", "north", "south", "row"],
        "Syllogisms & Relations": ["all", "some", "related", "conclusion", "photograph"],
    },
    "Verbal": {
        "Grammar": ["grammar", "grammatically", "fill in the blank", "sentence", "verb"],
        "Vocabulary": ["synonym", "antonym", "meaning", "word closest"],
        "Usage & Idioms": ["idiom", "usage", "punctuation", "correctly spelled"],
    },
    "Verbal Ability": {
        "Grammar": ["grammar", "grammatically", "fill in the blank", "sentence", "verb"],
        "Vocabulary": ["synonym", "antonym", "meaning", "word closest"],
        "Usage & Idioms": ["idiom", "usage", "punctuation", "correctly spelled"],
    },
    "DSA Basics": {
        "Data Structures": ["stack", "queue", "heap", "hash", "tree"],
        "Algorithms": ["sort", "search", "merge", "dfs", "bfs"],
        "Complexity": ["complexity", "worst-case", "average-case", "o("],
    },
    "SQL": {
        "Joins & Aggregation": ["join", "group", "having", "count", "sum", "max"],
        "DDL & DML": ["alter", "update", "delete", "truncate", "insert"],
        "Filtering & Keys": ["where", "like", "primary key", "distinct", "filter"],
    },
}


def infer_topic_category(topic_name, question_text):
    normalized = (question_text or "").lower()
    topic_patterns = TOPIC_CATEGORY_PATTERNS.get(topic_name, {})
    if not topic_patterns and topic_name in {"Logical Reasoning", "Verbal Ability"}:
        fallback_topic = "Logical" if topic_name == "Logical Reasoning" else "Verbal"
        topic_patterns = TOPIC_CATEGORY_PATTERNS.get(fallback_topic, {})
    for category, patterns in topic_patterns.items():
        if any(pattern in normalized for pattern in patterns):
            return category
    return topic_name


def build_test_performance_feedback(test_label, topic_name, questions, answers, recommendation_pool=None):
    category_stats = {}
    difficulty_stats = {}
    incorrect_categories = []
    correct_count = 0

    for question in questions:
        question_id = str(question.get("question_id", question.get("id", "")))
        selected_answer = (answers.get(question_id) or answers.get(question.get("question_id")) or answers.get(str(question.get("id"))) or "").upper()
        correct_answer = (question.get("correct_answer") or "").upper()
        is_correct = selected_answer == correct_answer
        if is_correct:
            correct_count += 1

        difficulty = question.get("difficulty", "Medium")
        question_topic = question.get("section") or topic_name
        category = infer_topic_category(question_topic, question.get("question_text", ""))

        category_bucket = category_stats.setdefault(category, {"correct": 0, "total": 0})
        category_bucket["total"] += 1
        if is_correct:
            category_bucket["correct"] += 1
        else:
            incorrect_categories.append(category)

        difficulty_bucket = difficulty_stats.setdefault(difficulty, {"correct": 0, "total": 0})
        difficulty_bucket["total"] += 1
        if is_correct:
            difficulty_bucket["correct"] += 1

    recommendations = []
    for category, stats in category_stats.items():
        accuracy = (stats["correct"] / stats["total"] * 100) if stats["total"] else 0
        if accuracy < 50:
            recommendations.append(f"Your accuracy in {category} questions is low.")

    easy_accuracy = (difficulty_stats.get("Easy", {}).get("correct", 0) / max(difficulty_stats.get("Easy", {}).get("total", 1), 1)) * 100 if difficulty_stats.get("Easy") else None
    medium_accuracy = (difficulty_stats.get("Medium", {}).get("correct", 0) / max(difficulty_stats.get("Medium", {}).get("total", 1), 1)) * 100 if difficulty_stats.get("Medium") else None
    hard_accuracy = (difficulty_stats.get("Hard", {}).get("correct", 0) / max(difficulty_stats.get("Hard", {}).get("total", 1), 1)) * 100 if difficulty_stats.get("Hard") else None

    if medium_accuracy is not None and easy_accuracy is not None and medium_accuracy + 15 < easy_accuracy:
        recommendations.append("Performance drops significantly in Medium-difficulty questions.")
    if hard_accuracy is not None and ((medium_accuracy is not None and hard_accuracy + 15 < medium_accuracy) or hard_accuracy < 40):
        recommendations.append("Hard-difficulty questions are reducing your score; revise fundamentals before moving to advanced sets.")

    repeated_mistakes = [category for category, count in {item: incorrect_categories.count(item) for item in set(incorrect_categories)}.items() if count >= 2]
    if repeated_mistakes:
        recommendations.append(f"Repeated mistakes are appearing in {', '.join(repeated_mistakes[:3])}. Practice those patterns before reattempting similar tests.")

    if recommendation_pool:
        recommendations.append(f"Recommended next tests based on weak areas: {', '.join(recommendation_pool[:3])}.")

    if not recommendations:
        recommendations.append(f"Good job on {test_label}. Keep building consistency and move to the next test set for stronger placement readiness.")

    difficulty_breakdown = {
        level: round((stats["correct"] / stats["total"] * 100), 2)
        for level, stats in difficulty_stats.items()
        if stats["total"]
    }
    category_breakdown = {
        category: round((stats["correct"] / stats["total"] * 100), 2)
        for category, stats in category_stats.items()
        if stats["total"]
    }

    return {
        "summary": f"You answered {correct_count} out of {len(questions)} correctly in {test_label}.",
        "recommendations": recommendations[:5],
        "category_breakdown": category_breakdown,
        "difficulty_breakdown": difficulty_breakdown,
        "repeated_mistakes": repeated_mistakes[:5],
    }
