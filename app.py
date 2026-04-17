# PLACIFY - AI-Powered Placement Readiness Platform
# Flask Backend Server

from flask import Flask, render_template, request, jsonify, session, redirect, url_for, send_file
import sqlite3
########
import sqlite3

conn = sqlite3.connect("database.db")
cursor = conn.cursor()

cursor.execute("SELECT * FROM users")
print(cursor.fetchall())

#########

import json
from pathlib import Path
from datetime import datetime
import re
############################################################
from datetime import datetime
from zoneinfo import ZoneInfo   # built-in

def current_time():
    return datetime.now(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S")
############################################################
import os 
from werkzeug.utils import secure_filename
from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

from data.test_catalog import COMPANY_TEST_SEED, TOPIC_TEST_CATALOG
from nlp.matcher import SkillMatcher
from nlp.resume_parser import ResumeParser
from nlp.scorer import ResumeScorer
from nlp.skill_extractor import SkillExtractor
from nlp.text_preprocessing import TextPreprocessor







app = Flask(__name__)
app.secret_key = 'placify-secret-key-change-in-production'

# Database configuration
DATABASE = 'database.db'
UPLOAD_FOLDER = 'uploads'
DOWNLOAD_FOLDER = 'downloads'
ALLOWED_RESUME_EXTENSIONS = {'.pdf', '.docx'}
ALLOWED_JD_EXTENSIONS = {'.pdf', '.docx', '.doc'}
DEFAULT_JOB_DESCRIPTION = (
    "We are hiring a software developer with python, java, javascript, sql, git, "
    "react, flask, docker, machine learning, and problem solving skills."
)
OPENAI_MODEL = os.environ.get('OPENAI_MODEL', 'gpt-5.2')

# Initialize database
def init_db():
    """Initialize database with schema"""
    if not os.path.exists(DATABASE):
        conn = sqlite3.connect(DATABASE)
        with open('sql/schema.sql', 'r') as f:
            conn.executescript(f.read())
        conn.commit()
        conn.close()
        print("Database initialized successfully!")
    ensure_resume_columns()
    ensure_topic_test_tables()
    ensure_company_test_questions()

# Database helper functions
def get_db():
    """Get database connection"""
    conn = sqlite3.connect(DATABASE)
    conn.row_factory = sqlite3.Row
    return conn

def query_db(query, args=(), one=False):
    """Execute SELECT query"""
    conn = get_db()
    cur = conn.execute(query, args)
    rv = cur.fetchall()
    conn.close()
    return (rv[0] if rv else None) if one else rv

def execute_db(query, args=()):
    """Execute INSERT/UPDATE/DELETE query"""
    conn = get_db()
    cur = conn.execute(query, args)
    conn.commit()
    last_id = cur.lastrowid
    conn.close()
    return last_id


def ensure_resume_columns():
    """Add resume columns needed by the current builder if they are missing."""
    conn = get_db()
    existing_columns = {
        row['name'] for row in conn.execute("PRAGMA table_info(resumes)").fetchall()
    }
    required_columns = {
        'target_company': "ALTER TABLE resumes ADD COLUMN target_company TEXT",
        'target_role': "ALTER TABLE resumes ADD COLUMN target_role TEXT",
        'job_description': "ALTER TABLE resumes ADD COLUMN job_description TEXT",
    }

    for column, statement in required_columns.items():
        if column not in existing_columns:
            conn.execute(statement)

    conn.commit()
    conn.close()


def ensure_topic_test_tables():
    """Create additive tables for fixed topic-wise tests."""
    conn = get_db()
    conn.executescript(
        '''
        CREATE TABLE IF NOT EXISTS topic_test_attempts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            test_key TEXT NOT NULL,
            topic_name TEXT NOT NULL,
            test_name TEXT NOT NULL,
            score DECIMAL(5,2) DEFAULT 0,
            total_questions INTEGER NOT NULL,
            correct_answers INTEGER DEFAULT 0,
            time_taken INTEGER DEFAULT 0,
            completed_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(user_id, test_key),
            FOREIGN KEY (user_id) REFERENCES users(id)
        );

        CREATE TABLE IF NOT EXISTS topic_test_responses (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            attempt_id INTEGER NOT NULL,
            question_key TEXT NOT NULL,
            selected_answer TEXT,
            is_correct BOOLEAN DEFAULT 0,
            FOREIGN KEY (attempt_id) REFERENCES topic_test_attempts(id)
        );
        '''
    )
    conn.commit()
    conn.close()


def ensure_company_test_questions():
    """Seed realistic company-wise question sets if the bank is empty."""
    conn = get_db()
    existing_count = conn.execute('SELECT COUNT(*) FROM company_test_questions').fetchone()[0]
    if existing_count:
        conn.close()
        return

    company_tests = conn.execute(
        '''
        SELECT ct.id, ct.test_name, c.company_name
        FROM company_tests ct
        JOIN companies c ON ct.company_id = c.id
        '''
    ).fetchall()

    for row in company_tests:
        company_key = (row['company_name'] or '').strip().lower()
        seed = COMPANY_TEST_SEED.get(company_key)
        if not seed:
            continue

        for section, question_text, options, correct_answer, explanation in seed['questions']:
            conn.execute(
                '''
                INSERT INTO company_test_questions (
                    company_test_id, section, question_type, question_text,
                    option_a, option_b, option_c, option_d, correct_answer, coding_output, points
                ) VALUES (?, ?, 'mcq', ?, ?, ?, ?, ?, ?, ?, 1)
                ''',
                [
                    row['id'],
                    section,
                    question_text,
                    options['A'],
                    options['B'],
                    options['C'],
                    options['D'],
                    correct_answer,
                    explanation,
                ]
            )

    conn.commit()
    conn.close()


def get_topic_test_map():
    test_map = {}
    for topic in TOPIC_TEST_CATALOG:
        for test in topic['tests']:
            test_map[test['test_id']] = test
    return test_map


TOPIC_TEST_MAP = get_topic_test_map()

# ========== ROUTES ==========

@app.route('/')
def index():
    """Landing page"""
    return render_template('index.html')

@app.route('/login')
def login_page():
    """Login page"""
    return render_template('login.html')

@app.route('/student')
def student_page():
    """Student dashboard page"""
    if 'user_id' not in session or session.get('role') != 'student':
        return redirect(url_for('login_page'))
    return render_template('student.html')

@app.route('/admin')
def admin_page():
    """Admin dashboard page"""
    if 'user_id' not in session or session.get('role') != 'admin':
        return redirect(url_for('login_page'))
    return render_template('admin.html')

@app.route('/tests')
def tests_page():
    """Tests page"""
    if 'user_id' not in session:
        return redirect(url_for('login_page'))
    return render_template('tests.html')

@app.route('/dashboard')
def dashboard_page():
    """Dashboard page"""
    if 'user_id' not in session:
        return redirect(url_for('login_page'))
    return render_template('dashboard.html')

@app.route('/resume')
def resume_page():
    """Resume builder page"""
    if 'user_id' not in session:
        return redirect(url_for('login_page'))
    return render_template('resume.html')


@app.route('/resume-analyzer')
def resume_analyzer_page():
    """Resume analyzer page"""
    if 'user_id' not in session:
        return redirect(url_for('login_page'))
    return render_template('resume_analyzer.html')

# ========== API ENDPOINTS ==========

@app.route('/api/login', methods=['POST'])
def api_login():
    """Login API"""
    data = request.get_json()
    username = data.get('username')
    password = data.get('password')
    
    user = query_db('SELECT * FROM users WHERE username = ? AND password = ?',
                    [username, password], one=True)
    
    if user:
        session['user_id'] = user['id']
        session['username'] = user['username']
        session['role'] = user['role']
        session['full_name'] = user['full_name']
        
        return jsonify({
            'success': True,
            'role': user['role'],
            'redirect': '/student' if user['role'] == 'student' else '/admin'
        })
    else:
        return jsonify({'success': False, 'message': 'Invalid credentials'}), 401

@app.route('/api/register', methods=['POST'])
def api_register():
    """Register new student"""
    data = request.get_json()
    
    try:
        execute_db('''INSERT INTO users (username, email, password, full_name, department, year, college)
                      VALUES (?, ?, ?, ?, ?, ?, ?)''',
                   [data.get('username'), data.get('email'), data.get('password'),
                    data.get('full_name'), data.get('department'), data.get('year'),
                    data.get('college', 'Placify College')])
        return jsonify({'success': True, 'message': 'Registration successful'})
    except sqlite3.IntegrityError:
        return jsonify({'success': False, 'message': 'Username or email already exists'}), 400

@app.route('/api/logout', methods=['POST'])
def api_logout():
    """Logout API"""
    session.clear()
    return jsonify({'success': True})

@app.route('/company-tests')
def company_tests_page():
    if 'user_id' not in session:
        return redirect(url_for('login_page'))
    return render_template('company_tests.html')


@app.route('/api/topic_tests', methods=['GET'])
def api_topic_tests():
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401

    attempts = query_db(
        '''
        SELECT test_key, score, correct_answers, total_questions, completed_at
        FROM topic_test_attempts
        WHERE user_id = ?
        ''',
        [session['user_id']]
    )
    attempt_map = {row['test_key']: dict(row) for row in attempts}

    payload = []
    for topic in TOPIC_TEST_CATALOG:
        topic_payload = {
            'topic_name': topic['topic_name'],
            'description': topic['description'],
            'test_count': topic['test_count'],
            'tests': [],
        }
        for test in topic['tests']:
            attempt = attempt_map.get(test['test_id'])
            topic_payload['tests'].append({
                'test_id': test['test_id'],
                'test_name': test['test_name'],
                'description': test['description'],
                'question_count': len(test['questions']),
                'time_limit': test['time_limit'],
                'attempted': bool(attempt),
                'status': 'Attempted' if attempt else 'Not Attempted',
                'score': attempt['score'] if attempt else None,
                'completed_at': attempt['completed_at'] if attempt else None,
            })
        payload.append(topic_payload)

    return jsonify(payload)


@app.route('/api/topic_tests/<test_id>/questions', methods=['GET'])
def api_topic_test_questions(test_id):
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401

    test = TOPIC_TEST_MAP.get(test_id)
    if not test:
        return jsonify({'success': False, 'message': 'Test not found'}), 404

    attempted = query_db(
        'SELECT id FROM topic_test_attempts WHERE user_id = ? AND test_key = ?',
        [session['user_id'], test_id],
        one=True
    )
    if attempted:
        return jsonify({'success': False, 'message': 'This test has already been attempted.'}), 400

    return jsonify({
        'test_id': test['test_id'],
        'test_name': test['test_name'],
        'topic_name': test['topic_name'],
        'time_limit': test['time_limit'],
        'questions': [
            {
                'question_id': question['question_id'],
                'question_text': question['question_text'],
                'option_a': question['options']['A'],
                'option_b': question['options']['B'],
                'option_c': question['options']['C'],
                'option_d': question['options']['D'],
            }
            for question in test['questions']
        ],
    })


@app.route('/api/topic_tests/<test_id>/submit', methods=['POST'])
def api_submit_topic_test(test_id):
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401

    test = TOPIC_TEST_MAP.get(test_id)
    if not test:
        return jsonify({'success': False, 'message': 'Test not found'}), 404

    existing_attempt = query_db(
        'SELECT id FROM topic_test_attempts WHERE user_id = ? AND test_key = ?',
        [session['user_id'], test_id],
        one=True
    )
    if existing_attempt:
        return jsonify({'success': False, 'message': 'Each topic test can only be attempted once.'}), 400

    data = request.get_json() or {}
    answers = data.get('answers', {})
    time_taken = data.get('time_taken', 0)
    correct_count = 0

    attempt_id = execute_db(
        '''
        INSERT INTO topic_test_attempts
        (user_id, test_key, topic_name, test_name, total_questions, time_taken, completed_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        ''',
        [session['user_id'], test['test_id'], test['topic_name'], test['test_name'], len(test['questions']), time_taken, current_time()]
    )

    for question in test['questions']:
        selected_answer = (answers.get(question['question_id']) or '').upper()
        is_correct = selected_answer == question['correct_answer']
        if is_correct:
            correct_count += 1

        execute_db(
            '''
            INSERT INTO topic_test_responses (attempt_id, question_key, selected_answer, is_correct)
            VALUES (?, ?, ?, ?)
            ''',
            [attempt_id, question['question_key'], selected_answer, is_correct]
        )

    score = round((correct_count / len(test['questions'])) * 100, 2) if test['questions'] else 0
    execute_db(
        'UPDATE topic_test_attempts SET score = ?, correct_answers = ? WHERE id = ?',
        [score, correct_count, attempt_id]
    )

    generate_ai_recommendations(session['user_id'])

    return jsonify({
        'success': True,
        'score': score,
        'correct': correct_count,
        'total': len(test['questions']),
        'attempt_id': attempt_id,
        'test_name': test['test_name'],
        'topic_name': test['topic_name'],
    })


@app.route('/api/topic_tests/<test_id>/solutions', methods=['GET'])
def api_topic_test_solutions(test_id):
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401

    test = TOPIC_TEST_MAP.get(test_id)
    if not test:
        return jsonify({'success': False, 'message': 'Test not found'}), 404

    attempt = query_db(
        'SELECT id, score, correct_answers, completed_at FROM topic_test_attempts WHERE user_id = ? AND test_key = ?',
        [session['user_id'], test_id],
        one=True
    )
    if not attempt:
        return jsonify({'success': False, 'message': 'Solutions are available only after attempting the test.'}), 403

    responses = query_db(
        'SELECT question_key, selected_answer, is_correct FROM topic_test_responses WHERE attempt_id = ?',
        [attempt['id']]
    )
    response_map = {row['question_key']: dict(row) for row in responses}

    return jsonify({
        'test_id': test['test_id'],
        'test_name': test['test_name'],
        'topic_name': test['topic_name'],
        'score': attempt['score'],
        'correct_answers': attempt['correct_answers'],
        'completed_at': attempt['completed_at'],
        'questions': [
            {
                'question_text': question['question_text'],
                'options': question['options'],
                'correct_answer': question['correct_answer'],
                'selected_answer': response_map.get(question['question_key'], {}).get('selected_answer', ''),
                'is_correct': bool(response_map.get(question['question_key'], {}).get('is_correct', 0)),
                'explanation': question['explanation'],
            }
            for question in test['questions']
        ],
    })

@app.route('/api/test_sections', methods=['GET'])
def api_test_sections():
    """Get all test sections"""
    sections = query_db('SELECT * FROM test_sections')
    return jsonify([dict(row) for row in sections])

@app.route('/api/questions/<int:section_id>', methods=['GET'])
def api_questions(section_id):
    """Get questions for a section"""
    questions = query_db('SELECT id, question_text, option_a, option_b, option_c, option_d FROM questions WHERE section_id = ?',
                         [section_id])
    return jsonify([dict(row) for row in questions])



@app.route('/api/submit_test', methods=['POST'])
def api_submit_test():
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401

    data = request.get_json()
    section_id = data.get('section_id')
    answers = data.get('answers', {})
    time_taken = data.get('time_taken', 0)  #  FIXED
    timestamp = current_time()               #  IST TIME

    questions = query_db(
        'SELECT id, correct_answer, points FROM questions WHERE section_id = ?',
        [section_id]
    )

    correct_count = 0
    total_points = 0
    earned_points = 0

    # ✅ Insert attempt WITH timestamp
    attempt_id = execute_db(
        '''INSERT INTO test_attempts 
           (user_id, section_id, total_questions, time_taken, completed_at)
           VALUES (?, ?, ?, ?, ?)''',
        [session['user_id'], section_id, len(questions), time_taken, timestamp]
    )

    for question in questions:
        q_id = question['id']
        correct_ans = question['correct_answer']
        points = question['points']
        total_points += points

        selected_ans = answers.get(str(q_id), '')
        is_correct = selected_ans.upper() == correct_ans.upper()

        if is_correct:
            correct_count += 1
            earned_points += points

        execute_db(
            '''INSERT INTO user_responses 
               (attempt_id, question_id, selected_answer, is_correct)
               VALUES (?, ?, ?, ?)''',
            [attempt_id, q_id, selected_ans, is_correct]
        )

    score = (earned_points / total_points * 100) if total_points > 0 else 0

    execute_db(
        'UPDATE test_attempts SET score = ?, correct_answers = ? WHERE id = ?',
        [score, correct_count, attempt_id]
    )

    generate_ai_recommendations(session['user_id'])

    return jsonify({
        'success': True,
        'score': round(score, 2),
        'correct': correct_count,
        'total': len(questions),
        'attempt_id': attempt_id
    })

    
@app.route('/api/company_tests', methods=['GET'])
def api_company_tests():
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401

    tests = query_db(
        """
        SELECT
            ct.id,
            c.company_name,
            ct.test_name,
            ct.total_duration,
            COALESCE(COUNT(ctq.id), ct.total_questions) AS total_questions,
            cta.score,
            cta.completed_at,
            CASE WHEN cta.id IS NULL THEN 0 ELSE 1 END AS attempted
        FROM company_tests ct
        JOIN companies c ON ct.company_id = c.id
        LEFT JOIN company_test_questions ctq ON ctq.company_test_id = ct.id
        LEFT JOIN company_test_attempts cta
            ON cta.company_test_id = ct.id AND cta.user_id = ?
        GROUP BY ct.id, c.company_name, ct.test_name, ct.total_duration, ct.total_questions, cta.id
        ORDER BY c.company_name
        """,
        [session['user_id']]
    )
    return jsonify([dict(row) for row in tests])


@app.route('/api/company_tests/<int:company_test_id>/questions', methods=['GET'])
def api_company_test_questions(company_test_id):
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401

    existing_attempt = query_db(
        'SELECT id FROM company_test_attempts WHERE user_id = ? AND company_test_id = ?',
        [session['user_id'], company_test_id],
        one=True
    )
    if existing_attempt:
        return jsonify({'success': False, 'message': 'This company-wise test has already been attempted.'}), 400

    test = query_db(
        '''
        SELECT ct.id, ct.test_name, ct.total_duration, c.company_name
        FROM company_tests ct
        JOIN companies c ON ct.company_id = c.id
        WHERE ct.id = ?
        ''',
        [company_test_id],
        one=True
    )
    questions = query_db(
        '''
        SELECT id, section, question_text, option_a, option_b, option_c, option_d
        FROM company_test_questions
        WHERE company_test_id = ?
        ORDER BY id
        ''',
        [company_test_id]
    )
    if not test or not questions:
        return jsonify({'success': False, 'message': 'Questions not available for this company test yet.'}), 404

    return jsonify({
        'company_test_id': company_test_id,
        'company_name': test['company_name'],
        'test_name': test['test_name'],
        'time_limit': test['total_duration'],
        'questions': [dict(row) for row in questions],
    })


@app.route('/api/company_tests/<int:company_test_id>/submit', methods=['POST'])
def api_submit_company_test(company_test_id):
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401

    existing_attempt = query_db(
        'SELECT id FROM company_test_attempts WHERE user_id = ? AND company_test_id = ?',
        [session['user_id'], company_test_id],
        one=True
    )
    if existing_attempt:
        return jsonify({'success': False, 'message': 'Each company-wise test can only be attempted once.'}), 400

    questions = query_db(
        '''
        SELECT id, correct_answer, points
        FROM company_test_questions
        WHERE company_test_id = ?
        ''',
        [company_test_id]
    )
    if not questions:
        return jsonify({'success': False, 'message': 'Questions not available for this company test yet.'}), 404

    data = request.get_json() or {}
    answers = data.get('answers', {})
    time_taken = data.get('time_taken', 0)
    correct_count = 0
    total_points = 0

    attempt_id = execute_db(
        '''
        INSERT INTO company_test_attempts (user_id, company_test_id, time_taken, completed_at)
        VALUES (?, ?, ?, ?)
        ''',
        [session['user_id'], company_test_id, time_taken, current_time()]
    )

    for question in questions:
        selected_answer = (answers.get(str(question['id'])) or '').upper()
        is_correct = selected_answer == (question['correct_answer'] or '').upper()
        total_points += question['points'] or 1
        if is_correct:
            correct_count += 1

        execute_db(
            '''
            INSERT INTO company_test_responses (attempt_id, question_id, selected_answer, is_correct)
            VALUES (?, ?, ?, ?)
            ''',
            [attempt_id, question['id'], selected_answer, is_correct]
        )

    score = round((correct_count / len(questions)) * 100, 2) if questions else 0
    execute_db(
        'UPDATE company_test_attempts SET score = ? WHERE id = ?',
        [score, attempt_id]
    )

    return jsonify({
        'success': True,
        'score': score,
        'correct': correct_count,
        'total': len(questions),
        'attempt_id': attempt_id,
    })


@app.route('/api/company_tests/<int:company_test_id>/solutions', methods=['GET'])
def api_company_test_solutions(company_test_id):
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401

    attempt = query_db(
        'SELECT id, score, completed_at FROM company_test_attempts WHERE user_id = ? AND company_test_id = ?',
        [session['user_id'], company_test_id],
        one=True
    )
    if not attempt:
        return jsonify({'success': False, 'message': 'Solutions are available only after attempting the company-wise test.'}), 403

    test = query_db(
        '''
        SELECT ct.test_name, c.company_name
        FROM company_tests ct
        JOIN companies c ON c.id = ct.company_id
        WHERE ct.id = ?
        ''',
        [company_test_id],
        one=True
    )
    questions = query_db(
        '''
        SELECT id, section, question_text, option_a, option_b, option_c, option_d, correct_answer, coding_output
        FROM company_test_questions
        WHERE company_test_id = ?
        ORDER BY id
        ''',
        [company_test_id]
    )
    responses = query_db(
        '''
        SELECT question_id, selected_answer, is_correct
        FROM company_test_responses
        WHERE attempt_id = ?
        ''',
        [attempt['id']]
    )
    response_map = {row['question_id']: dict(row) for row in responses}

    return jsonify({
        'company_name': test['company_name'],
        'test_name': test['test_name'],
        'score': attempt['score'],
        'completed_at': attempt['completed_at'],
        'questions': [
            {
                'section': row['section'],
                'question_text': row['question_text'],
                'options': {
                    'A': row['option_a'],
                    'B': row['option_b'],
                    'C': row['option_c'],
                    'D': row['option_d'],
                },
                'correct_answer': row['correct_answer'],
                'selected_answer': response_map.get(row['id'], {}).get('selected_answer', ''),
                'is_correct': bool(response_map.get(row['id'], {}).get('is_correct', 0)),
                'explanation': row['coding_output'] or 'Review the correct option and compare it with your choice.',
            }
            for row in questions
        ]
    })

@app.route('/api/user_scores', methods=['GET'])
def api_user_scores():
    """Get user's test scores"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401
    
    scores = query_db(
        '''
        SELECT ta.score, ta.correct_answers, ta.total_questions, ta.completed_at, ts.section_name
        FROM test_attempts ta
        JOIN test_sections ts ON ta.section_id = ts.id
        WHERE ta.user_id = ?
        UNION ALL
        SELECT tta.score, tta.correct_answers, tta.total_questions, tta.completed_at,
               tta.topic_name || ' - ' || tta.test_name AS section_name
        FROM topic_test_attempts tta
        WHERE tta.user_id = ?
        ORDER BY completed_at DESC
        ''',
        [session['user_id'], session['user_id']]
    )
    
    return jsonify([dict(row) for row in scores])

@app.route('/api/section_performance', methods=['GET'])
def api_section_performance():
    """Get section-wise performance"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401
    
    performance = query_db(
        '''
        SELECT section_name, AVG(score) as avg_score, COUNT(*) as attempts
        FROM (
            SELECT ts.section_name AS section_name, ta.score AS score
            FROM test_attempts ta
            JOIN test_sections ts ON ta.section_id = ts.id
            WHERE ta.user_id = ?
            UNION ALL
            SELECT tta.topic_name AS section_name, tta.score AS score
            FROM topic_test_attempts tta
            WHERE tta.user_id = ?
        ) combined_scores
        GROUP BY section_name
        ''',
        [session['user_id'], session['user_id']]
    )
    
    return jsonify([dict(row) for row in performance])

from flask import request, jsonify


@app.route('/api/save_resume', methods=['POST'])
def api_save_resume():
    """Save resume and calculate score"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401

    data = request.form.to_dict() if request.form else (request.get_json() or {})
    user_id = session['user_id']
    job_description_text = (data.get('job_description') or '').strip()
    jd_file = request.files.get('jd_file')

    if jd_file and jd_file.filename:
        jd_text, jd_extension = extract_uploaded_document_text(jd_file, ALLOWED_JD_EXTENSIONS)
        if jd_extension not in ALLOWED_JD_EXTENSIONS:
            return jsonify({'success': False, 'message': 'Only PDF, DOCX, and DOC files are allowed for the job description'}), 400
        job_description_text = jd_text.strip()

    tailored_resume = build_tailored_resume(data, job_description_text)
    resume_text = tailored_resume['resume_text']
    scores = calculate_resume_score(data, tailored_resume['match_percentage'])
    download_paths = generate_resume_files(user_id, tailored_resume['sections'])
    
    # Check if resume exists
    existing = query_db('SELECT id FROM resumes WHERE user_id = ?', [user_id], one=True)
    
    if existing:
        # Update existing resume
        execute_db('''UPDATE resumes SET full_name=?, email=?, phone=?, education=?, skills=?,
                      experience=?, projects=?, certifications=?, target_company=?, target_role=?,
                      job_description=?, resume_text=?,
                      ats_score=?, keyword_score=?, format_score=?, overall_score=?, feedback=?,
                      updated_at=CURRENT_TIMESTAMP
                      WHERE user_id=?''',
                   [data.get('full_name'), data.get('email'), data.get('phone'),
                    data.get('education'), data.get('skills'), data.get('experience'),
                    data.get('projects'), data.get('certifications'),
                    data.get('target_company'), data.get('target_role'), job_description_text, resume_text,
                    scores['ats_score'], scores['keyword_score'], scores['format_score'],
                    scores['overall_score'], scores['feedback'], user_id])
    else:
        # Insert new resume
        execute_db('''INSERT INTO resumes (user_id, full_name, email, phone, education, skills,
                      experience, projects, certifications, target_company, target_role,
                      job_description, resume_text, ats_score, keyword_score, format_score,
                      overall_score, feedback)
                      VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                   [user_id, data.get('full_name'), data.get('email'), data.get('phone'),
                    data.get('education'), data.get('skills'), data.get('experience'),
                    data.get('projects'), data.get('certifications'),
                    data.get('target_company'), data.get('target_role'), job_description_text, resume_text,
                    scores['ats_score'], scores['keyword_score'], scores['format_score'],
                    scores['overall_score'], scores['feedback']])

    return jsonify({
        'success': True,
        'scores': scores,
        'resume_text': resume_text,
        'download_pdf_url': '/api/download_resume/pdf',
        'download_doc_url': '/api/download_resume/doc',
    })

@app.route('/api/get_resume', methods=['GET'])
def api_get_resume():
    """Get user's resume"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401
    
    resume = query_db('SELECT * FROM resumes WHERE user_id = ?', [session['user_id']], one=True)
    
    if resume:
        return jsonify(dict(resume))
    else:
        return jsonify({'success': False, 'message': 'No resume found'}), 404


@app.route('/api/download_resume/<file_type>', methods=['GET'])
def api_download_resume(file_type):
    """Download generated resume files."""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401

    base_path = Path(DOWNLOAD_FOLDER)
    if file_type == 'pdf':
        file_path = base_path / f"resume_{session['user_id']}.pdf"
        download_name = "placify_resume.pdf"
    elif file_type == 'doc':
        file_path = base_path / f"resume_{session['user_id']}.docx"
        download_name = "placify_resume.docx"
    else:
        return jsonify({'success': False, 'message': 'Invalid file type'}), 400

    if not file_path.exists():
        return jsonify({'success': False, 'message': 'Resume file not found'}), 404

    return send_file(file_path, as_attachment=True, download_name=download_name)


@app.route('/api/analyze_resume', methods=['POST'])
def api_analyze_resume():
    """Analyze uploaded resume using the NLP pipeline"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401

    resume_file = request.files.get('resume_file')
    if not resume_file or not resume_file.filename:
        return jsonify({'success': False, 'message': 'Resume file is required'}), 400

    extracted_text, file_extension = extract_uploaded_document_text(resume_file, ALLOWED_RESUME_EXTENSIONS)
    if file_extension not in ALLOWED_RESUME_EXTENSIONS:
        return jsonify({'success': False, 'message': 'Only PDF and DOCX files are allowed'}), 400

    job_description_text = (request.form.get('job_description') or '').strip()
    jd_file = request.files.get('jd_file')
    if jd_file and jd_file.filename:
        jd_text, jd_extension = extract_uploaded_document_text(jd_file, ALLOWED_JD_EXTENSIONS)
        if jd_extension not in ALLOWED_JD_EXTENSIONS:
            return jsonify({'success': False, 'message': 'Only PDF, DOCX, and DOC files are allowed for the job description'}), 400
        job_description_text = jd_text.strip()

    if not job_description_text:
        job_description_text = DEFAULT_JOB_DESCRIPTION

    preprocessor = TextPreprocessor()
    extractor = SkillExtractor()
    scorer = ResumeScorer()

    cleaned_resume_text = preprocessor.preprocess(extracted_text)
    cleaned_job_text = preprocessor.preprocess(job_description_text)

    resume_skills = extractor.extract_skills(cleaned_resume_text)
    job_skills = extractor.extract_skills(cleaned_job_text)
    match_result = SkillMatcher.match(resume_skills, job_skills)
    normalized_resume_skills = {skill.strip().lower() for skill in resume_skills}
    normalized_job_skills = {skill.strip().lower() for skill in job_skills}
    missing_skills = sorted(normalized_job_skills - normalized_resume_skills)
    extra_skills = sorted(normalized_resume_skills - normalized_job_skills)
    keyword_overlap = sorted(normalized_resume_skills & normalized_job_skills)
    resume_score = scorer.calculate_score(
        match_result['match_percentage'],
        count_projects(extracted_text),
    )
    resume_type_result = detect_resume_type(extracted_text)
    capability_scores = compute_capability_match_scores(
        extracted_text,
        job_description_text,
        resume_skills,
        job_skills,
        resume_type_result['resume_type'],
    )
    skill_confidence = score_skills_with_evidence(resume_skills, extracted_text)

    return jsonify({
        'success': True,
        'extracted_skills': resume_skills,
        'skill_confidence': skill_confidence,
        'matched_skills': match_result['matched_skills'],
        'skill_match_percentage': match_result['match_percentage'],
        'match_percentage': match_result['match_percentage'],
        'missing_skills': missing_skills,
        'extra_skills': extra_skills,
        'keyword_overlap': keyword_overlap,
        'resume_text': extracted_text,
        'job_description_text': job_description_text,
        'resume_type': resume_type_result['resume_type'],
        'resume_type_confidence': resume_type_result['confidence_score'],
        'resume_capabilities': capability_scores['resume_capabilities'],
        'jd_capabilities': capability_scores['jd_capabilities'],
        'capability_category_scores': capability_scores['category_scores'],
        'role_fit_score': capability_scores['role_fit_score'],
        'skill_readiness_score': capability_scores['skill_readiness_score'],
        'learning_gap_score': capability_scores['learning_gap_score'],
        'resume_score': resume_score,
    })


@app.route('/api/resume_ai_suggestions', methods=['POST'])
def api_resume_ai_suggestions():
    """Generate AI-powered resume improvement suggestions."""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401

    data = request.get_json() or {}
    resume_text = (data.get('resume_text') or '').strip()
    job_description_text = (data.get('job_description_text') or '').strip()

    if not resume_text or not job_description_text:
        return jsonify({'success': False, 'message': 'Resume text and job description are required'}), 400

    try:
        suggestions = generate_resume_ai_suggestions(resume_text, job_description_text)
    except ValueError as error:
        return jsonify({'success': False, 'message': str(error)}), 400
    except Exception as error:
        return jsonify({'success': False, 'message': f'AI suggestion generation failed: {error}'}), 500

    return jsonify({'success': True, 'suggestions': suggestions})


@app.route('/api/resume_ai_chat', methods=['POST'])
def api_resume_ai_chat():
    """Answer follow-up questions about resume improvement."""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401

    data = request.get_json() or {}
    resume_text = (data.get('resume_text') or '').strip()
    job_description_text = (data.get('job_description_text') or '').strip()
    question = (data.get('question') or '').strip()
    suggestions = data.get('suggestions') or {}
    chat_history = data.get('chat_history') or []

    if not resume_text or not job_description_text or not question:
        return jsonify({'success': False, 'message': 'Resume text, job description, and question are required'}), 400

    try:
        answer = generate_resume_ai_chat_reply(
            resume_text,
            job_description_text,
            suggestions,
            question,
            chat_history,
        )
    except ValueError as error:
        return jsonify({'success': False, 'message': str(error)}), 400
    except Exception as error:
        return jsonify({'success': False, 'message': f'AI chat failed: {error}'}), 500

    return jsonify({'success': True, 'answer': answer})

@app.route('/api/ai_recommendations', methods=['GET'])
def api_ai_recommendations():
    """Get AI-based recommendations"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401
    
    recommendations = query_db('SELECT * FROM ai_recommendations WHERE user_id = ? ORDER BY generated_at DESC LIMIT 1',
                               [session['user_id']], one=True)
    
    if recommendations:
        return jsonify(dict(recommendations))
    else:
        # Generate if not exists
        generate_ai_recommendations(session['user_id'])
        recommendations = query_db('SELECT * FROM ai_recommendations WHERE user_id = ? ORDER BY generated_at DESC LIMIT 1',
                                   [session['user_id']], one=True)
        return jsonify(dict(recommendations) if recommendations else {})

@app.route('/api/admin/students', methods=['GET'])
def api_admin_students():
    """Get all students (admin only)"""
    if 'user_id' not in session or session.get('role') != 'admin':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    
    students = query_db('''SELECT u.*, 
                          AVG(ta.score) as avg_score,
                          COUNT(DISTINCT ta.section_id) as sections_attempted
                          FROM users u
                          LEFT JOIN test_attempts ta ON u.id = ta.user_id
                          WHERE u.role = 'student'
                          GROUP BY u.id''')
    
    return jsonify([dict(row) for row in students])

@app.route('/api/admin/department_stats', methods=['GET'])
def api_admin_department_stats():
    """Get department-wise statistics (admin only)"""
    if 'user_id' not in session or session.get('role') != 'admin':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    
    stats = query_db('''SELECT u.department, 
                       COUNT(DISTINCT u.id) as student_count,
                       AVG(ta.score) as avg_score,
                       COUNT(ta.id) as total_attempts
                       FROM users u
                       LEFT JOIN test_attempts ta ON u.id = ta.user_id
                       WHERE u.role = 'student' AND u.department IS NOT NULL
                       GROUP BY u.department''')
    
    return jsonify([dict(row) for row in stats])


@app.route('/api/admin/reset_demo_data', methods=['POST'])
def api_admin_reset_demo_data():
    """Clear user-generated demo data while preserving system users."""
    if 'user_id' not in session or session.get('role') != 'admin':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    try:
        deleted_rows = reset_demo_data()
    except sqlite3.DatabaseError as error:
        return jsonify({'success': False, 'message': f'Database reset failed: {error}'}), 500

    return jsonify({
        'success': True,
        'message': 'Demo data cleared successfully',
        'deleted_rows': deleted_rows,
    })

@app.route('/api/user_info', methods=['GET'])
def api_user_info():
    """Get current user info"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401
    
    return jsonify({
        'user_id': session['user_id'],
        'username': session['username'],
        'full_name': session['full_name'],
        'role': session['role']
    })

# ========== HELPER FUNCTIONS ==========

def calculate_resume_score(resume_data, jd_match_percentage=0):
    """Calculate resume scores using rule-based AI"""
    
    # ATS Score - Check for key sections
    ats_score = 0
    if resume_data.get('full_name'): ats_score += 15
    if resume_data.get('email'): ats_score += 10
    if resume_data.get('phone'): ats_score += 10
    if resume_data.get('education'): ats_score += 20
    if resume_data.get('skills'): ats_score += 20
    if resume_data.get('experience') or resume_data.get('projects'): ats_score += 25
    
    # Keyword Score - Check for important keywords
    keyword_score = 0
    skills_text = (resume_data.get('skills', '') + ' ' + 
                   resume_data.get('experience', '') + ' ' + 
                   resume_data.get('projects', '')).lower()
    
    important_keywords = ['python', 'java', 'javascript', 'react', 'sql', 'database',
                         'api', 'git', 'team', 'project', 'leadership', 'communication',
                         'problem solving', 'agile', 'development']
    
    for keyword in important_keywords:
        if keyword in skills_text:
            keyword_score += 100 / len(important_keywords)
    
    # Format Score - Check text length and structure
    format_score = 0
    total_length = sum([len(str(resume_data.get(key, ''))) for key in 
                       ['education', 'skills', 'experience', 'projects', 'certifications']])
    
    if total_length > 100: format_score += 40
    if total_length > 300: format_score += 30
    if total_length > 500: format_score += 30
    
    # Overall score
    overall_score = (ats_score + keyword_score + format_score) / 3
    overall_score = (overall_score * 0.6) + (jd_match_percentage * 0.4)
    
    # Generate feedback
    feedback = []
    if ats_score < 70:
        feedback.append("Add more details to key sections like education and experience.")
    if keyword_score < 50:
        feedback.append("Include more relevant technical skills and keywords.")
    if format_score < 60:
        feedback.append("Expand your resume with more detailed descriptions.")
    if overall_score >= 80:
        feedback.append("Excellent resume! Well structured and comprehensive.")
    elif overall_score >= 60:
        feedback.append("Good resume, but there's room for improvement.")
    else:
        feedback.append("Your resume needs significant enhancement.")
    
    return {
        'ats_score': round(ats_score, 2),
        'keyword_score': round(keyword_score, 2),
        'format_score': round(format_score, 2),
        'overall_score': round(overall_score, 2),
        'feedback': ' '.join(feedback),
        'jd_match_percentage': round(jd_match_percentage, 2),
    }

def generate_ai_recommendations(user_id):
    """Generate AI-based recommendations for user"""
    
    # Get user's performance
    performance = query_db(
        '''
        SELECT section_name, AVG(score) as avg_score
        FROM (
            SELECT ts.section_name AS section_name, ta.score AS score
            FROM test_attempts ta
            JOIN test_sections ts ON ta.section_id = ts.id
            WHERE ta.user_id = ?
            UNION ALL
            SELECT tta.topic_name AS section_name, tta.score AS score
            FROM topic_test_attempts tta
            WHERE tta.user_id = ?
        ) grouped_scores
        GROUP BY section_name
        ''',
        [user_id, user_id]
    )
    
    weak_sections = []
    improvement_areas = []
    overall_avg = 0
    
    if performance:
        scores = [row['avg_score'] for row in performance]
        overall_avg = sum(scores) / len(scores) if scores else 0
        
        # Identify weak sections (below 60%)
        for row in performance:
            if row['avg_score'] < 60:
                weak_sections.append(row['section_name'])
                
                # Add specific improvement suggestions
                if row['section_name'] == 'Aptitude':
                    improvement_areas.append("Practice more quantitative problems and speed calculations")
                elif row['section_name'] in ('Logical Reasoning', 'Logical'):
                    improvement_areas.append("Work on pattern recognition and logical puzzles")
                elif row['section_name'] in ('Coding', 'DSA Basics'):
                    improvement_areas.append("Focus on data structures and algorithms")
                elif row['section_name'] == 'Verbal':
                    improvement_areas.append("Improve vocabulary, grammar accuracy, and reading comprehension speed")
                elif row['section_name'] == 'SQL':
                    improvement_areas.append("Revise joins, grouping, filtering, and query writing fundamentals")
                elif row['section_name'] == 'HR & Soft Skills':
                    improvement_areas.append("Improve communication and behavioral interview skills")
                elif row['section_name'] == 'Domain Knowledge':
                    improvement_areas.append("Study core technical concepts and fundamentals")
    
    # Practice focus recommendations
    if overall_avg < 50:
        practice_focus = "Focus on fundamentals across all sections. Take more practice tests."
    elif overall_avg < 70:
        practice_focus = "Good progress! Concentrate on weak areas and time management."
    else:
        practice_focus = "Excellent performance! Maintain consistency and polish advanced topics."
    
    # Readiness score
    readiness_score = min(overall_avg + 10, 100)  # Bonus points for attempt
    
    # Store recommendations
    execute_db('''INSERT INTO ai_recommendations 
                  (user_id, weak_sections, improvement_areas, practice_focus, readiness_score)
                  VALUES (?, ?, ?, ?, ?)''',
               [user_id, json.dumps(weak_sections), json.dumps(improvement_areas),
                practice_focus, readiness_score])
    
    return True


def reset_demo_data():
    """Delete user-generated rows while preserving predefined users."""
    conn = get_db()
    deleted_rows = {}

    try:
        conn.execute('BEGIN')

        protected_user_query = """
            SELECT id
            FROM users
            WHERE username IN (?, ?)
        """
        protected_users = ('admin', 'student1')

        cursor = conn.execute(
            f"""
            DELETE FROM user_responses
            WHERE attempt_id IN (
                SELECT id
                FROM test_attempts
                WHERE user_id NOT IN ({protected_user_query})
            )
            """,
            protected_users,
        )
        deleted_rows['user_responses'] = cursor.rowcount

        cursor = conn.execute(
            f"""
            DELETE FROM company_test_responses
            WHERE attempt_id IN (
                SELECT id
                FROM company_test_attempts
                WHERE user_id NOT IN ({protected_user_query})
            )
            """,
            protected_users,
        )
        deleted_rows['company_test_responses'] = cursor.rowcount

        cursor = conn.execute(
            f"""
            DELETE FROM test_attempts
            WHERE user_id NOT IN ({protected_user_query})
            """,
            protected_users,
        )
        deleted_rows['test_attempts'] = cursor.rowcount

        cursor = conn.execute(
            f"""
            DELETE FROM company_test_attempts
            WHERE user_id NOT IN ({protected_user_query})
            """,
            protected_users,
        )
        deleted_rows['company_test_attempts'] = cursor.rowcount

        cursor = conn.execute(
            f"""
            DELETE FROM resumes
            WHERE user_id NOT IN ({protected_user_query})
            """,
            protected_users,
        )
        deleted_rows['resumes'] = cursor.rowcount

        cursor = conn.execute(
            f"""
            DELETE FROM ai_recommendations
            WHERE user_id NOT IN ({protected_user_query})
            """,
            protected_users,
        )
        deleted_rows['ai_recommendations'] = cursor.rowcount

        conn.commit()
        return deleted_rows
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


def count_projects(text):
    """Estimate project count from resume text."""
    matches = re.findall(r"\bproject[s]?\b", text.lower())
    return len(matches)


def detect_resume_type(resume_text):
    """Detect whether a resume is for a fresher or experienced professional."""
    text = (resume_text or '').lower()
    project_count = count_projects(text)
    experience_entries = len(re.findall(r"\bexperience\b|\bworked\b|\bemployment\b|\bdeveloper\b", text))
    internship_hits = len(re.findall(r"\binternship\b|\bintern\b|\btrainee\b", text))

    years = []
    years.extend(float(match) for match in re.findall(r"(\d+(?:\.\d+)?)\+?\s+years?", text))
    years.extend(float(match) for match in re.findall(r"experience\s+of\s+(\d+(?:\.\d+)?)", text))
    max_years = max(years) if years else 0.0

    fresher_score = 0.0
    experienced_score = 0.0

    if max_years >= 2:
        experienced_score += 0.55
    elif max_years >= 1:
        experienced_score += 0.35
    else:
        fresher_score += 0.25

    if internship_hits:
        fresher_score += min(0.25, internship_hits * 0.08)

    if project_count >= max(experience_entries, 1):
        fresher_score += 0.25
    else:
        experienced_score += 0.2

    if experience_entries >= 2:
        experienced_score += 0.25
    elif experience_entries == 1:
        experienced_score += 0.1
    else:
        fresher_score += 0.15

    total_score = fresher_score + experienced_score
    if total_score == 0:
        return {'resume_type': 'Fresher', 'confidence_score': 0.5}

    fresher_confidence = fresher_score / total_score
    experienced_confidence = experienced_score / total_score

    if experienced_confidence > fresher_confidence:
        return {
            'resume_type': 'Experienced professional',
            'confidence_score': round(experienced_confidence, 2),
        }

    return {
        'resume_type': 'Fresher',
        'confidence_score': round(fresher_confidence, 2),
    }


def split_items(text):
    """Split comma or newline separated text into clean items."""
    return [
        item.strip()
        for item in re.split(r'[\n,]+', text or '')
        if item.strip()
    ]


def extract_capabilities(text, skills=None):
    """Extract capability buckets from resume or JD text."""
    normalized_text = (text or '').lower()
    extracted_skills = list(skills or [])

    technical_keywords = {
        'api', 'backend', 'frontend', 'database', 'algorithms', 'data structures',
        'machine learning', 'deep learning', 'testing', 'automation', 'cloud',
    }
    domain_keywords = {
        'finance', 'banking', 'healthcare', 'ecommerce', 'education', 'analytics',
        'nlp', 'computer vision', 'cybersecurity', 'devops', 'web development',
    }
    tool_keywords = {
        'git', 'github', 'docker', 'kubernetes', 'aws', 'azure', 'google cloud',
        'flask', 'django', 'fastapi', 'react', 'node.js', 'express.js',
        'tensorflow', 'pytorch', 'pandas', 'numpy', 'scikit-learn',
    }

    technical_skills = sorted({
        skill for skill in extracted_skills
        if skill in technical_keywords or skill in tool_keywords
    } | {keyword for keyword in technical_keywords if keyword in normalized_text})
    domain_knowledge = sorted({
        skill for skill in extracted_skills if skill in domain_keywords
    } | {keyword for keyword in domain_keywords if keyword in normalized_text})
    tools_and_technologies = sorted({
        skill for skill in extracted_skills if skill in tool_keywords
    } | {keyword for keyword in tool_keywords if keyword in normalized_text})

    applied_work = []
    for label, patterns in {
        'projects': [r'\bproject[s]?\b', r'\bbuilt\b', r'\bdeveloped\b', r'\bcreated\b'],
        'experience': [r'\bexperience\b', r'\bworked\b', r'\bemployment\b', r'\bengineer\b'],
        'internships': [r'\binternship\b', r'\bintern\b', r'\btrainee\b'],
    }.items():
        if any(re.search(pattern, normalized_text) for pattern in patterns):
            applied_work.append(label)

    return {
        'technical_skills': technical_skills,
        'applied_work': applied_work,
        'domain_knowledge': domain_knowledge,
        'tools_and_technologies': tools_and_technologies,
    }


def compute_overlap_score(resume_values, jd_values):
    """Compute a normalized overlap score between resume and JD values."""
    resume_set = {value.strip().lower() for value in resume_values if value and value.strip()}
    jd_set = {value.strip().lower() for value in jd_values if value and value.strip()}

    if not jd_set:
        return 100.0

    return round((len(resume_set & jd_set) / len(jd_set)) * 100, 2)


def get_capability_weights(resume_type):
    """Get dynamic capability weights based on resume type."""
    if resume_type == 'Experienced professional':
        return {
            'technical_skills': 0.3,
            'applied_work': 0.35,
            'domain_knowledge': 0.2,
            'tools_and_technologies': 0.15,
        }

    return {
        'technical_skills': 0.35,
        'applied_work': 0.2,
        'domain_knowledge': 0.15,
        'tools_and_technologies': 0.3,
    }


def compute_capability_match_scores(resume_text, job_description_text, resume_skills, job_skills, resume_type):
    """Compute capability-based role fit metrics."""
    resume_capabilities = extract_capabilities(resume_text, resume_skills)
    jd_capabilities = extract_capabilities(job_description_text, job_skills)
    weights = get_capability_weights(resume_type)

    category_scores = {
        key: compute_overlap_score(resume_capabilities[key], jd_capabilities[key])
        for key in weights
    }

    role_fit_score = round(
        sum(category_scores[key] * weights[key] for key in weights),
        2,
    )
    skill_readiness_score = round(
        (category_scores['technical_skills'] * 0.6) +
        (category_scores['tools_and_technologies'] * 0.25) +
        (category_scores['domain_knowledge'] * 0.15),
        2,
    )
    learning_gap_score = round(max(0.0, 100.0 - role_fit_score), 2)

    return {
        'resume_capabilities': resume_capabilities,
        'jd_capabilities': jd_capabilities,
        'category_scores': category_scores,
        'role_fit_score': role_fit_score,
        'skill_readiness_score': skill_readiness_score,
        'learning_gap_score': learning_gap_score,
    }


def infer_skill_evidence(skill, resume_text):
    """Find supporting evidence for a skill from projects, experience, and tools usage."""
    normalized_text = (resume_text or '').lower()
    normalized_skill = skill.strip().lower()
    evidence = {
        'projects': [],
        'experience': [],
        'tools_usage': [],
    }

    project_patterns = [
        rf"project[s]?[^\n.]*{re.escape(normalized_skill)}",
        rf"{re.escape(normalized_skill)}[^\n.]*project[s]?",
    ]
    experience_patterns = [
        rf"experience[^\n.]*{re.escape(normalized_skill)}",
        rf"worked[^\n.]*{re.escape(normalized_skill)}",
        rf"{re.escape(normalized_skill)}[^\n.]*experience",
    ]
    tool_patterns = [
        rf"using[^\n.]*{re.escape(normalized_skill)}",
        rf"with[^\n.]*{re.escape(normalized_skill)}",
        rf"tool[s]?[^\n.]*{re.escape(normalized_skill)}",
    ]

    for pattern in project_patterns:
        evidence['projects'].extend(re.findall(pattern, normalized_text))
    for pattern in experience_patterns:
        evidence['experience'].extend(re.findall(pattern, normalized_text))
    for pattern in tool_patterns:
        evidence['tools_usage'].extend(re.findall(pattern, normalized_text))

    return {key: list(dict.fromkeys(value)) for key, value in evidence.items()}


def get_skill_confidence_label(score):
    """Map a numeric evidence score to a confidence label."""
    if score >= 75:
        return 'High'
    if score >= 45:
        return 'Medium'
    return 'Low'


def score_skills_with_evidence(skills, resume_text):
    """Score extracted skills based on supporting evidence in the resume."""
    skill_confidence = []

    for skill in skills:
        evidence = infer_skill_evidence(skill, resume_text)
        score = 20.0
        if evidence['projects']:
            score += 35.0
        if evidence['experience']:
            score += 35.0
        if evidence['tools_usage']:
            score += 20.0
        if not any(evidence.values()):
            score -= 15.0

        score = max(5.0, min(score, 100.0))
        skill_confidence.append({
            'skill': skill,
            'confidence_score': round(score, 2),
            'confidence_level': get_skill_confidence_label(score),
            'evidence': evidence,
        })

    return skill_confidence


def build_tailored_resume(data, job_description_text):
    """Create ATS-friendly resume content tailored to the provided job description."""
    preprocessor = TextPreprocessor()
    extractor = SkillExtractor()

    skills_input = data.get('skills', '')
    experience = data.get('experience', '')
    projects = data.get('projects', '')
    certifications = data.get('certifications', '')
    education = data.get('education', '')
    target_company = (data.get('target_company') or 'the target company').strip()
    target_role = (data.get('target_role') or 'the target role').strip()

    job_description_text = (job_description_text or DEFAULT_JOB_DESCRIPTION).strip()
    combined_user_text = " ".join([skills_input, experience, projects, certifications, education])

    user_skills = extractor.extract_skills(preprocessor.preprocess(combined_user_text))
    job_skills = extractor.extract_skills(preprocessor.preprocess(job_description_text))
    merged_skills = list(dict.fromkeys(user_skills + job_skills))
    match_percentage = 100.0 if job_skills else 90.0
    ats_keywords = ", ".join(job_skills) if job_skills else "python, sql, communication"

    summary = (
        f"Targeting the {target_role} opportunity at {target_company}. "
        f"Candidate brings experience across {', '.join(merged_skills[:8]) or 'software development'} "
        f"with a resume aligned to the job description requirements."
    )

    sections = [
        ("Full Name", data.get('full_name', '')),
        ("Contact", f"Email: {data.get('email', '')}\nPhone: {data.get('phone', '')}"),
        ("Professional Summary", summary),
        ("Core Skills", ", ".join(merged_skills)),
        ("Professional Experience", experience or "Add relevant internships, work experience, or leadership roles."),
        ("Projects", projects or "Add project details with measurable outcomes and tools used."),
        ("Education", education),
        ("Certifications", certifications or "Add certifications relevant to the target role."),
        ("Target Company", target_company),
        ("Target Role", target_role),
        ("Job Description Alignment", job_description_text),
        ("ATS Keywords", ats_keywords),
    ]

    resume_text = "\n\n".join(f"{title}\n{content}".strip() for title, content in sections if content)
    return {
        'sections': sections,
        'resume_text': resume_text,
        'match_percentage': match_percentage,
    }


def generate_resume_files(user_id, sections):
    """Generate downloadable PDF and DOCX resume files."""
    output_dir = Path(DOWNLOAD_FOLDER)
    output_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = output_dir / f"resume_{user_id}.pdf"
    docx_path = output_dir / f"resume_{user_id}.docx"

    generate_resume_pdf(pdf_path, sections)
    generate_resume_docx(docx_path, sections)

    return {
        'pdf': str(pdf_path),
        'doc': str(docx_path),
    }


def generate_resume_pdf(file_path, sections):
    """Create a simple ATS-friendly PDF resume."""
    pdf = canvas.Canvas(str(file_path), pagesize=LETTER)
    width, height = LETTER
    x_margin = 50
    y = height - 50

    for title, content in sections:
        if not content:
            continue

        if y < 100:
            pdf.showPage()
            y = height - 50

        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawString(x_margin, y, title)
        y -= 18

        pdf.setFont("Helvetica", 10)
        for line in str(content).splitlines():
            wrapped_lines = wrap_text(line, 90)
            for wrapped_line in wrapped_lines:
                if y < 60:
                    pdf.showPage()
                    y = height - 50
                    pdf.setFont("Helvetica", 10)
                pdf.drawString(x_margin, y, wrapped_line)
                y -= 14
        y -= 10

    pdf.save()


def generate_resume_docx(file_path, sections):
    """Create a simple ATS-friendly DOCX resume."""
    document = Document()

    for index, (title, content) in enumerate(sections):
        if not content:
            continue
        if index == 0:
            document.add_heading(str(content), level=0)
            continue
        document.add_heading(title, level=1)
        for line in str(content).splitlines():
            document.add_paragraph(line)

    document.save(str(file_path))


def wrap_text(text, max_length):
    """Wrap text into lines suitable for PDF rendering."""
    words = text.split()
    lines = []
    current_line = []

    for word in words:
        trial_line = " ".join(current_line + [word])
        if len(trial_line) <= max_length:
            current_line.append(word)
        else:
            lines.append(" ".join(current_line))
            current_line = [word]

    if current_line:
        lines.append(" ".join(current_line))

    return lines or [""]


def save_uploaded_file(uploaded_file):
    """Save an uploaded file into the uploads folder."""
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    filename = secure_filename(uploaded_file.filename)
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    uploaded_file.save(file_path)
    return file_path


def extract_uploaded_document_text(uploaded_file, allowed_extensions):
    """Extract text from an uploaded document."""
    file_extension = os.path.splitext(uploaded_file.filename)[1].lower()
    if file_extension not in allowed_extensions:
        return None, file_extension

    file_path = save_uploaded_file(uploaded_file)

    if file_extension == '.doc':
        with open(file_path, 'rb') as file:
            return file.read().decode('latin-1', errors='ignore'), file_extension

    return ResumeParser.extract_text(file_path), file_extension


def get_openai_client():
    """Create an OpenAI client from environment configuration."""
    api_key = os.environ.get('OPENAI_API_KEY')
    if not api_key:
        raise ValueError('OPENAI_API_KEY is not set')

    try:
        from openai import OpenAI
    except ImportError as error:
        raise ValueError('OpenAI SDK is not installed') from error

    return OpenAI(api_key=api_key)


def extract_json_object(text):
    """Extract and parse a JSON object from model output."""
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        start = text.find('{')
        end = text.rfind('}')
        if start == -1 or end == -1 or end <= start:
            raise ValueError('AI response was not valid JSON')
        return json.loads(text[start:end + 1])


def normalize_suggestions_payload(payload):
    """Normalize the AI suggestion payload into a predictable shape."""
    suggestions = payload if isinstance(payload, dict) else {}
    section_improvements = suggestions.get('section_wise_improvements') or []

    return {
        'skill_gaps': suggestions.get('skill_gaps') or [],
        'missing_keywords': suggestions.get('missing_keywords') or [],
        'section_wise_improvements': section_improvements,
        'ats_optimization_tips': suggestions.get('ats_optimization_tips') or [],
        'shortlisting_probability_estimate': suggestions.get('shortlisting_probability_estimate') or 'Unknown',
    }


def generate_resume_ai_suggestions(resume_text, job_description_text):
    """Generate structured AI suggestions for a resume against a job description."""
    client = get_openai_client()
    prompt = (
        "Return only valid JSON with these keys: "
        "skill_gaps (array of strings), "
        "missing_keywords (array of strings), "
        "section_wise_improvements (array of objects with section and improvement), "
        "ats_optimization_tips (array of strings), "
        "shortlisting_probability_estimate (string)."
    )
    response = client.responses.create(
        model=OPENAI_MODEL,
        instructions=(
            "You are an expert ATS resume reviewer. Compare the resume against the job description. "
            "Be practical, concise, and specific."
        ),
        input=[
            {
                'role': 'user',
                'content': [
                    {
                        'type': 'input_text',
                        'text': (
                            f"{prompt}\n\nResume Text:\n{resume_text}\n\n"
                            f"Job Description Text:\n{job_description_text}"
                        ),
                    }
                ],
            }
        ],
    )
    return normalize_suggestions_payload(extract_json_object(response.output_text))


def generate_resume_ai_chat_reply(resume_text, job_description_text, suggestions, question, chat_history):
    """Generate a follow-up chat response using the analyzed resume context."""
    client = get_openai_client()
    messages = [
        {
            'role': 'user',
            'content': [
                {
                    'type': 'input_text',
                    'text': (
                        "Use the following context to answer follow-up resume improvement questions.\n\n"
                        f"Resume Text:\n{resume_text}\n\n"
                        f"Job Description Text:\n{job_description_text}\n\n"
                        f"Current Suggestions:\n{json.dumps(suggestions)}"
                    ),
                }
            ],
        }
    ]

    for message in chat_history[-6:]:
        role = message.get('role')
        content = (message.get('content') or '').strip()
        if role in {'user', 'assistant'} and content:
            messages.append(
                {
                    'role': role,
                    'content': [{'type': 'input_text', 'text': content}],
                }
            )

    messages.append(
        {
            'role': 'user',
            'content': [{'type': 'input_text', 'text': question}],
        }
    )

    response = client.responses.create(
        model=OPENAI_MODEL,
        instructions=(
            "You are a resume coach. Answer follow-up questions using the supplied resume, "
            "job description, and prior suggestions. Keep responses clear and actionable."
        ),
        input=messages,
    )
    return response.output_text.strip()


def run_resume_analysis_demo():
    """Demonstrate the NLP resume analysis pipeline and print outputs."""
    resume_file = os.environ.get("RESUME_FILE")
    job_description = os.environ.get(
        "JOB_DESCRIPTION",
        "Looking for a Python developer with Flask, SQL, Git, Docker, and React experience.",
    )

    if resume_file and os.path.exists(resume_file):
        extracted_text = ResumeParser.extract_text(resume_file)
        print(f"Resume source: {resume_file}")
    else:
        extracted_text = (
            "John Doe\n"
            "Python Flask SQL Git Docker React developer\n"
            "Project: Resume Analyzer using NLP and Machine Learning\n"
            "Project: Placement Portal built with Flask and PostgreSQL"
        )
        print("Resume source: sample text (set RESUME_FILE to analyze a PDF or DOCX)")

    preprocessor = TextPreprocessor()
    extractor = SkillExtractor()
    scorer = ResumeScorer()

    cleaned_resume_text = preprocessor.preprocess(extracted_text)
    cleaned_job_text = preprocessor.preprocess(job_description)

    resume_skills = extractor.extract_skills(cleaned_resume_text)
    job_skills = extractor.extract_skills(cleaned_job_text)
    match_result = SkillMatcher.match(resume_skills, job_skills)
    project_count = count_projects(extracted_text)
    resume_score = scorer.calculate_score(match_result["match_percentage"], project_count)

    print("\nResume Analysis Pipeline")
    print("=" * 50)
    print("Extracted Text:")
    print(extracted_text)
    print("\nCleaned Resume Text:")
    print(cleaned_resume_text)
    print("\nResume Skills:")
    print(resume_skills)
    print("\nJob Description Skills:")
    print(job_skills)
    print("\nMatched Skills:")
    print(match_result["matched_skills"])
    print(f"Match Percentage: {match_result['match_percentage']}%")
    print(f"Project Count: {project_count}")
    print(f"Resume Score: {resume_score}/100")
    print("=" * 50)

# ========== MAIN ==========

if __name__ == '__main__':
    init_db()
    run_resume_analysis_demo()
    print("=" * 50)
    print("PLACIFY - Placement Readiness Platform")
    print("=" * 50)
    print("\nServer starting at http://127.0.0.1:5000")
    print("\nDefault Login Credentials:")
    print("Admin - Username: admin, Password: admin123")
    print("Student - Username: student1, Password: student123")
    print("=" * 50)
    app.run(debug=True, host='0.0.0.0', port=5000)
